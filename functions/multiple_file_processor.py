from fastapi import APIRouter, HTTPException
from pathlib import Path
from typing import List, Dict
import os
from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document as DocxDocument
from PIL import Image
import mimetypes
from fastapi import FastAPI
import uvicorn
from langchain_core.documents import Document
import os
from database.extracted_texts import upsert_upload_record

import google.generativeai as genai
import PIL.Image

from database.connection import database
extracted_texts_collection = database.get_collection("extracted_texts")


UPLOAD_DIR = Path("uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

def extract_text_from_txt(file_path: Path) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
        return file.read()

def extract_text_from_pdf(file_path: Path) -> str:
    try:
        return extract_pdf_text(file_path)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(file_path: Path) -> str:
    try:
        doc = DocxDocument(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract text from DOCX: {str(e)}")

def extract_text_from_image(file_path: Path) -> str:
        """Extracts text from an image (placeholder for future implementation)."""
        try:
            GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
            genai.configure(api_key=GOOGLE_API_KEY)
            model = genai.GenerativeModel("gemini-1.5-flash")
            image = PIL.Image.open(file_path)
            prompt= "for the given input image.if the image has text , extract the text from the image and return the text.if the image has only objects, return the summary of the image"
            response = model.generate_content([prompt, image])
            return response
        except Exception as e:
            raise ValueError(f"Error reading image file: {e}")
        


# Loader mapping for supported MIME types
LOADER_MAPPING = {
    "text/plain": extract_text_from_txt,
    "application/pdf": extract_text_from_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": extract_text_from_docx,
    "image/jpeg": extract_text_from_image,
    "image/png": extract_text_from_image,
}

async def process_folder(folder_path: Path) -> str:
    overall_text = ""

    for item in folder_path.iterdir():
        if item.is_dir():
            # Recursively process subfolders
            overall_text += await process_folder(item)
        elif item.is_file():
            mime_type, _ = mimetypes.guess_type(item)

            if mime_type in LOADER_MAPPING:
                extractor = LOADER_MAPPING[mime_type]
                text = extractor(item)
                overall_text += f"\n\n--- Extracted from {item.name} ---\n{text}"

            else:
                raise HTTPException(
                    status_code=415,
                    detail=f"Unsupported file type: {item.name} ({mime_type})",
                )

    return overall_text

async def process_files(session_id: str):
    documents=[]
    session_folder = UPLOAD_DIR / session_id

    if not session_folder.exists() or not session_folder.is_dir():
        raise HTTPException(status_code=404, detail="Session folder not found.")

    overall_text = await process_folder(session_folder)
    # print("overall_text extracted successfully")
    # # print(type(overall_text))
    
    if overall_text:
        # Convert overall text into a Document format
        documents.append(Document(page_content=overall_text, metadata={"title": session_id}))
        # Save the document content to MongoDB
        await upsert_upload_record(
            extracted_texts_collection, session_id, [{"content": overall_text, "metadata": {"title": session_id}}]
        )
    else:
        raise HTTPException(status_code=204, detail="No content extracted from files.")

    return documents
   